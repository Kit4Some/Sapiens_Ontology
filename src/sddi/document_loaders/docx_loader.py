"""
DOCX Document Loader.

Loads Microsoft Word (.docx) documents for the SDDI pipeline.
"""

import io
from pathlib import Path
from typing import BinaryIO, Any

import structlog

from src.sddi.document_loaders.base import BaseDocumentLoader
from src.sddi.state import RawDocument

logger = structlog.get_logger(__name__)


class DOCXLoader(BaseDocumentLoader):
    """
    Loader for Microsoft Word documents (.docx).

    Uses python-docx for text extraction with support for:
    - Paragraphs and headings
    - Tables (converted to markdown)
    - Lists (bulleted and numbered)
    - Headers and footers (optional)
    - Document metadata
    """

    SUPPORTED_EXTENSIONS = ["docx"]

    def __init__(
        self,
        include_headers_footers: bool = False,
        include_tables: bool = True,
        preserve_formatting: bool = True,
        one_doc_per_section: bool = False,
    ):
        """
        Initialize the DOCX loader.

        Args:
            include_headers_footers: If True, include header/footer text
            include_tables: If True, extract and convert tables to markdown
            preserve_formatting: If True, preserve headings and list formatting
            one_doc_per_section: If True, create one document per major section
        """
        self.include_headers_footers = include_headers_footers
        self.include_tables = include_tables
        self.preserve_formatting = preserve_formatting
        self.one_doc_per_section = one_doc_per_section

    async def load(
        self,
        file: BinaryIO,
        filename: str,
        metadata: dict[str, Any] | None = None,
    ) -> list[RawDocument]:
        """Load a DOCX file."""
        logger.info("Loading DOCX file", filename=filename)

        content = file.read()

        try:
            return await self._load_with_python_docx(content, filename, metadata)
        except ImportError:
            raise ImportError(
                "python-docx is required for DOCX loading. Install with:\n"
                "  pip install python-docx"
            )

    async def _load_with_python_docx(
        self,
        content: bytes,
        filename: str,
        metadata: dict[str, Any] | None = None,
    ) -> list[RawDocument]:
        """Load DOCX using python-docx library."""
        from docx import Document
        from docx.opc.exceptions import PackageNotFoundError

        try:
            doc = Document(io.BytesIO(content))
        except PackageNotFoundError:
            logger.error("Invalid DOCX file", filename=filename)
            raise ValueError(f"Invalid or corrupted DOCX file: {filename}")

        # Extract document properties
        doc_props = self._extract_document_properties(doc)

        logger.info(
            "DOCX opened",
            filename=filename,
            paragraphs=len(doc.paragraphs),
            tables=len(doc.tables),
            sections=len(doc.sections),
        )

        if self.one_doc_per_section:
            return self._extract_per_section(doc, filename, metadata, doc_props)

        # Extract all content
        all_text = []

        # Headers (optional)
        if self.include_headers_footers:
            header_text = self._extract_headers(doc)
            if header_text:
                all_text.append(f"--- HEADER ---\n{header_text}\n--- END HEADER ---\n")

        # Main content
        main_content = self._extract_paragraphs(doc)
        if main_content:
            all_text.append(main_content)

        # Tables
        if self.include_tables and doc.tables:
            tables_text = self._extract_tables(doc)
            if tables_text:
                all_text.append(tables_text)

        # Footers (optional)
        if self.include_headers_footers:
            footer_text = self._extract_footers(doc)
            if footer_text:
                all_text.append(f"--- FOOTER ---\n{footer_text}\n--- END FOOTER ---")

        combined_text = "\n\n".join(filter(None, all_text))

        if not combined_text.strip():
            logger.warning("No text extracted from DOCX", filename=filename)
            return []

        raw_doc = RawDocument(
            id=self._generate_doc_id(filename),
            content=combined_text,
            source=filename,
            metadata={
                "file_type": "docx",
                "paragraph_count": len(doc.paragraphs),
                "table_count": len(doc.tables),
                "section_count": len(doc.sections),
                "extractor": "python-docx",
                **doc_props,
                **(metadata or {}),
            },
        )

        logger.info(
            "DOCX loaded",
            filename=filename,
            char_count=len(combined_text),
            paragraphs=len(doc.paragraphs),
            tables=len(doc.tables),
        )

        return [raw_doc]

    def _extract_document_properties(self, doc) -> dict[str, Any]:
        """Extract document metadata/properties."""
        props = {}

        try:
            core_props = doc.core_properties
            if core_props.title:
                props["title"] = core_props.title
            if core_props.author:
                props["author"] = core_props.author
            if core_props.subject:
                props["subject"] = core_props.subject
            if core_props.keywords:
                props["keywords"] = core_props.keywords
            if core_props.created:
                props["created_at"] = core_props.created.isoformat() if core_props.created else None
            if core_props.modified:
                props["modified_at"] = core_props.modified.isoformat() if core_props.modified else None
        except Exception as e:
            logger.debug("Could not extract document properties", error=str(e))

        return props

    def _extract_paragraphs(self, doc) -> str:
        """Extract text from paragraphs with optional formatting preservation."""
        text_parts = []

        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue

            if self.preserve_formatting:
                # Check for headings
                if para.style and para.style.name:
                    style_name = para.style.name.lower()
                    if "heading 1" in style_name:
                        text = f"# {text}"
                    elif "heading 2" in style_name:
                        text = f"## {text}"
                    elif "heading 3" in style_name:
                        text = f"### {text}"
                    elif "heading 4" in style_name:
                        text = f"#### {text}"
                    elif "title" in style_name:
                        text = f"# {text}"
                    elif "list" in style_name:
                        text = f"- {text}"

            text_parts.append(text)

        return "\n\n".join(text_parts)

    def _extract_tables(self, doc) -> str:
        """Extract tables and convert to markdown format."""
        tables_text = []

        for table_idx, table in enumerate(doc.tables):
            table_rows = []

            for row_idx, row in enumerate(table.rows):
                cells = [cell.text.strip().replace("\n", " ") for cell in row.cells]
                table_rows.append("| " + " | ".join(cells) + " |")

                # Add header separator after first row
                if row_idx == 0:
                    separator = "| " + " | ".join(["---"] * len(cells)) + " |"
                    table_rows.append(separator)

            if table_rows:
                tables_text.append(f"[Table {table_idx + 1}]\n" + "\n".join(table_rows))

        return "\n\n".join(tables_text)

    def _extract_headers(self, doc) -> str:
        """Extract text from document headers."""
        header_texts = []

        for section in doc.sections:
            try:
                header = section.header
                if header:
                    for para in header.paragraphs:
                        text = para.text.strip()
                        if text:
                            header_texts.append(text)
            except Exception:
                pass

        return "\n".join(header_texts)

    def _extract_footers(self, doc) -> str:
        """Extract text from document footers."""
        footer_texts = []

        for section in doc.sections:
            try:
                footer = section.footer
                if footer:
                    for para in footer.paragraphs:
                        text = para.text.strip()
                        if text:
                            footer_texts.append(text)
            except Exception:
                pass

        return "\n".join(footer_texts)

    def _extract_per_section(
        self,
        doc,
        filename: str,
        metadata: dict[str, Any] | None,
        doc_props: dict[str, Any],
    ) -> list[RawDocument]:
        """Extract one document per major section (Heading 1)."""
        documents = []
        current_section_title = "Introduction"
        current_section_content: list[str] = []
        section_idx = 0

        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue

            # Check if this is a Heading 1 (new section)
            style_name = para.style.name.lower() if para.style else ""
            is_heading1 = "heading 1" in style_name or "title" in style_name

            if is_heading1 and current_section_content:
                # Save previous section
                section_text = "\n\n".join(current_section_content)
                if section_text.strip():
                    raw_doc = RawDocument(
                        id=self._generate_doc_id(filename, section_idx),
                        content=f"# {current_section_title}\n\n{section_text}",
                        source=filename,
                        metadata={
                            "file_type": "docx",
                            "section_index": section_idx,
                            "section_title": current_section_title,
                            "extractor": "python-docx",
                            **doc_props,
                            **(metadata or {}),
                        },
                    )
                    documents.append(raw_doc)
                    section_idx += 1

                # Start new section
                current_section_title = text
                current_section_content = []
            else:
                # Add to current section
                if self.preserve_formatting and "heading" in style_name:
                    level = self._get_heading_level(style_name)
                    text = "#" * level + " " + text
                current_section_content.append(text)

        # Don't forget the last section
        if current_section_content:
            section_text = "\n\n".join(current_section_content)
            if section_text.strip():
                raw_doc = RawDocument(
                    id=self._generate_doc_id(filename, section_idx),
                    content=f"# {current_section_title}\n\n{section_text}",
                    source=filename,
                    metadata={
                        "file_type": "docx",
                        "section_index": section_idx,
                        "section_title": current_section_title,
                        "extractor": "python-docx",
                        **doc_props,
                        **(metadata or {}),
                    },
                )
                documents.append(raw_doc)

        logger.info(
            "DOCX sections extracted",
            filename=filename,
            section_count=len(documents),
        )

        return documents

    def _get_heading_level(self, style_name: str) -> int:
        """Get heading level from style name."""
        style_lower = style_name.lower()
        for i in range(1, 7):
            if f"heading {i}" in style_lower:
                return i
        if "title" in style_lower:
            return 1
        return 2  # Default to h2


class DOCLoader(BaseDocumentLoader):
    """
    Loader for legacy Word documents (.doc).

    Note: .doc format requires additional dependencies.
    Consider converting to .docx for best results.
    """

    SUPPORTED_EXTENSIONS = ["doc"]

    async def load(
        self,
        file: BinaryIO,
        filename: str,
        metadata: dict[str, Any] | None = None,
    ) -> list[RawDocument]:
        """Load a DOC file."""
        logger.warning(
            "Legacy .doc format detected. For best results, convert to .docx",
            filename=filename,
        )

        content = file.read()

        # Try textract first, then fallback to antiword
        try:
            return await self._load_with_textract(content, filename, metadata)
        except ImportError:
            pass

        try:
            return await self._load_with_antiword(content, filename, metadata)
        except Exception as e:
            logger.error("Failed to load .doc file", filename=filename, error=str(e))
            raise ImportError(
                "Cannot load legacy .doc files. Options:\n"
                "  1. Convert to .docx using Microsoft Word or LibreOffice\n"
                "  2. Install textract: pip install textract\n"
                "  3. Install antiword: apt-get install antiword (Linux)"
            )

    async def _load_with_textract(
        self,
        content: bytes,
        filename: str,
        metadata: dict[str, Any] | None = None,
    ) -> list[RawDocument]:
        """Load DOC using textract."""
        import textract
        import tempfile
        import os

        # textract needs a file path
        with tempfile.NamedTemporaryFile(suffix=".doc", delete=False) as tmp:
            tmp.write(content)
            tmp_path = tmp.name

        try:
            text = textract.process(tmp_path).decode("utf-8")
        finally:
            os.unlink(tmp_path)

        if not text.strip():
            logger.warning("No text extracted from DOC", filename=filename)
            return []

        raw_doc = RawDocument(
            id=self._generate_doc_id(filename),
            content=text,
            source=filename,
            metadata={
                "file_type": "doc",
                "extractor": "textract",
                **(metadata or {}),
            },
        )

        return [raw_doc]

    async def _load_with_antiword(
        self,
        content: bytes,
        filename: str,
        metadata: dict[str, Any] | None = None,
    ) -> list[RawDocument]:
        """Load DOC using antiword (Linux only)."""
        import subprocess
        import tempfile
        import os

        with tempfile.NamedTemporaryFile(suffix=".doc", delete=False) as tmp:
            tmp.write(content)
            tmp_path = tmp.name

        try:
            result = subprocess.run(
                ["antiword", tmp_path],
                capture_output=True,
                text=True,
                timeout=60,
            )
            text = result.stdout
        finally:
            os.unlink(tmp_path)

        if not text.strip():
            logger.warning("No text extracted from DOC", filename=filename)
            return []

        raw_doc = RawDocument(
            id=self._generate_doc_id(filename),
            content=text,
            source=filename,
            metadata={
                "file_type": "doc",
                "extractor": "antiword",
                **(metadata or {}),
            },
        )

        return [raw_doc]
